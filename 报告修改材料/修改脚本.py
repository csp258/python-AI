"""
Script to modify the research report:
1. Generate chart images
2. Change title
3. Expand Section 4 with deep analysis
4. Add richer data tables
5. Insert charts into document
"""
import os
import shutil
from copy import deepcopy
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np

BASE = r'd:\下载\visual studio code\text'
ORIGINAL = os.path.join(BASE, '研究报告_AI赋能乡村教育.docx')
BACKUP = os.path.join(BASE, '研究报告_AI赋能乡村教育_backup.docx')

# ============================================================
# STEP 0: Backup original
# ============================================================
shutil.copy2(ORIGINAL, BACKUP)
print(f"Backup created: {BACKUP}")

# ============================================================
# STEP 1: Generate chart images
# ============================================================
plt.rcParams['font.family'] = 'sans-serif'
plt.rcParams['font.sans-serif'] = ['Arial', 'DejaVu Sans', 'Liberation Sans']
plt.rcParams['figure.dpi'] = 150
plt.rcParams['savefig.dpi'] = 150
plt.rcParams['savefig.bbox'] = 'tight'
plt.rcParams['savefig.facecolor'] = 'white'

# Color palette - professional, accessible
C1 = '#2c3e50'  # dark blue-grey
C2 = '#3498db'  # blue
C3 = '#2ecc71'  # green
C4 = '#e74c3c'  # red
C5 = '#f39c12'  # orange
C6 = '#9b59b6'  # purple
COLORS_3 = ['#3498db', '#2ecc71', '#e74c3c']
COLORS_4 = ['#3498db', '#2ecc71', '#f39c12', '#e74c3c']

# --- Chart 1: Cost-Effectiveness Comparison (Dual Bar) ---
fig, ax1 = plt.subplots(figsize=(8, 5))
models = ['Guizhou\n1:1 Tablets', 'Yunnan\nShared Classroom', 'Gansu\nSatellite+Offline']
cost_per_student = [1040, 342, 365]
gain_per_10k = [0.19, 0.63, 0.51]
x = np.arange(len(models))
width = 0.35
bars1 = ax1.bar(x - width/2, cost_per_student, width, color=COLORS_3[0], alpha=0.85, label='Cost per Student/Year (¥)', zorder=3)
ax1.set_ylabel('Cost per Student/Year (¥)', fontsize=11, fontweight='bold', color=COLORS_3[0])
ax1.tick_params(axis='y', labelcolor=COLORS_3[0])
ax2 = ax1.twinx()
bars2 = ax2.bar(x + width/2, gain_per_10k, width, color=COLORS_3[2], alpha=0.85, label='Learning Gain per ¥10K (points)', zorder=3)
ax2.set_ylabel('Learning Gain per ¥10K Invested (points)', fontsize=11, fontweight='bold', color=COLORS_3[2])
ax2.tick_params(axis='y', labelcolor=COLORS_3[2])
ax1.set_xticks(x)
ax1.set_xticklabels(models, fontsize=10)
ax1.set_title('Figure 6: Cost-Effectiveness Comparison of Three AI Deployment Models', fontsize=13, fontweight='bold', pad=15)
# Combined legend
bars = [bars1, bars2]
labels = ['Cost per Student/Year (¥)', 'Learning Gain per ¥10K (points)']
ax1.legend(bars, labels, loc='upper left', fontsize=9, framealpha=0.9)
ax1.grid(axis='y', alpha=0.3, zorder=0)
# Add value labels on bars
for bar in bars1:
    ax1.text(bar.get_x() + bar.get_width()/2., bar.get_height() + 12, f'¥{bar.get_height():.0f}', ha='center', fontsize=9, fontweight='bold')
for bar in bars2:
    ax2.text(bar.get_x() + bar.get_width()/2., bar.get_height() + 0.01, f'{bar.get_height():.2f}', ha='center', fontsize=9, fontweight='bold')
plt.tight_layout()
plt.savefig(os.path.join(BASE, 'chart_cost_effectiveness.png'))
plt.close()
print("Chart 1 saved.")

# --- Chart 2: Learning Gains by Baseline Quartile (Horizontal Bar) ---
fig, ax = plt.subplots(figsize=(8, 4.5))
quartiles = ['Q1 (Bottom 25%)', 'Q2 (25–50%)', 'Q3 (50–75%)', 'Q4 (Top 25%)']
gains = [6.1, 11.8, 14.7, 9.3]
colors = ['#e74c3c', '#f39c12', '#2ecc71', '#3498db']
bars = ax.barh(quartiles, gains, color=colors, alpha=0.85, height=0.55, zorder=3)
for bar, val in zip(bars, gains):
    ax.text(bar.get_width() + 0.3, bar.get_y() + bar.get_height()/2., f'+{val} pts', va='center', fontsize=11, fontweight='bold')
ax.set_xlabel('Mean Score Gain (points)', fontsize=11, fontweight='bold')
ax.set_title('Figure 7: Learning Gains by Baseline Performance Quartile', fontsize=13, fontweight='bold', pad=15)
ax.set_xlim(0, 19)
ax.grid(axis='x', alpha=0.3, zorder=0)
ax.spines['top'].set_visible(False)
ax.spines['right'].set_visible(False)
plt.tight_layout()
plt.savefig(os.path.join(BASE, 'chart_quartile_gains.png'))
plt.close()
print("Chart 2 saved.")

# --- Chart 3: Teacher Confidence Comparison (Grouped Bar) ---
fig, ax = plt.subplots(figsize=(8, 5))
categories = ['Very Confident', 'Confident', 'Somewhat Confident', 'Not Confident']
case_a = [8, 33, 42, 17]
case_b = [31, 45, 19, 5]
case_c = [5, 28, 44, 23]
x = np.arange(len(categories))
width = 0.25
ax.bar(x - width, case_a, width, color=COLORS_3[0], alpha=0.85, label='Case A: Guizhou', zorder=3)
ax.bar(x, case_b, width, color=COLORS_3[1], alpha=0.85, label='Case B: Yunnan', zorder=3)
ax.bar(x + width, case_c, width, color=COLORS_3[2], alpha=0.85, label='Case C: Gansu', zorder=3)
ax.set_xticks(x)
ax.set_xticklabels(categories, fontsize=10)
ax.set_ylabel('Percentage of Teachers (%)', fontsize=11, fontweight='bold')
ax.set_title('Figure 8: Teacher Confidence in AI-Integrated Teaching by Case (After 18 Months)', fontsize=13, fontweight='bold', pad=15)
ax.legend(fontsize=9, framealpha=0.9)
ax.grid(axis='y', alpha=0.3, zorder=0)
# Add value labels
for i, (ca, cb, cc) in enumerate(zip(case_a, case_b, case_c)):
    ax.text(i - width, ca + 0.8, f'{ca}%', ha='center', fontsize=8)
    ax.text(i, cb + 0.8, f'{cb}%', ha='center', fontsize=8)
    ax.text(i + width, cc + 0.8, f'{cc}%', ha='center', fontsize=8)
plt.tight_layout()
plt.savefig(os.path.join(BASE, 'chart_teacher_confidence.png'))
plt.close()
print("Chart 3 saved.")

# --- Chart 4: 5-Year TCO Breakdown by Category (Stacked Bar) ---
fig, ax = plt.subplots(figsize=(8, 5.5))
categories_tco = ['Hardware\n(Initial+Replace)', 'Software\nLicensing', 'Connectivity', 'Teacher\nTraining', 'Technical\nSupport']
guizhou_tco = [260.4, 120, 48, 90, 105.6]  # in ¥K
yunnan_tco = [62.4, 90, 48, 210, 101.6]
gansu_tco = [53.2, 60, 144, 72, 48.8]
x = np.arange(len(categories_tco))
width = 0.28
ax.bar(x - width, guizhou_tco, width, color=COLORS_3[0], alpha=0.85, label='Guizhou (1:1 Tablets) — ¥624K total', zorder=3)
ax.bar(x, yunnan_tco, width, color=COLORS_3[1], alpha=0.85, label='Yunnan (Shared Classroom) — ¥512K total', zorder=3)
ax.bar(x + width, gansu_tco, width, color=COLORS_3[2], alpha=0.85, label='Gansu (Satellite+Offline) — ¥438K total', zorder=3)
ax.set_xticks(x)
ax.set_xticklabels(categories_tco, fontsize=9)
ax.set_ylabel('5-Year Cost (¥ Thousands)', fontsize=11, fontweight='bold')
ax.set_title('Figure 9: Five-Year Total Cost of Ownership Breakdown by Category', fontsize=13, fontweight='bold', pad=15)
ax.legend(fontsize=8, framealpha=0.9, loc='upper right')
ax.grid(axis='y', alpha=0.3, zorder=0)
plt.tight_layout()
plt.savefig(os.path.join(BASE, 'chart_tco_breakdown.png'))
plt.close()
print("Chart 4 saved.")

# --- Chart 5: Implementation Phases — Projected Learning Gains Trajectory ---
fig, ax = plt.subplots(figsize=(8, 4.5))
phases = ['Baseline\n(2024)', 'Phase 1\nFoundation\n(Year 1–2)', 'Phase 2\nIntegration\n(Year 3–5)', 'Phase 3\nOptimization\n(Year 5–10)']
math_gains = [0, 8, 16, 24]
english_gains = [0, 7, 14, 21]
chinese_gains = [0, 3, 7, 12]
x = np.arange(len(phases))
ax.plot(x, math_gains, 'o-', color=COLORS_3[0], linewidth=2.5, markersize=8, label='Mathematics')
ax.plot(x, english_gains, 's-', color=COLORS_3[1], linewidth=2.5, markersize=8, label='English')
ax.plot(x, chinese_gains, '^--', color=COLORS_3[2], linewidth=2.5, markersize=8, label='Chinese')
ax.fill_between(x, 0, math_gains, alpha=0.08, color=COLORS_3[0])
ax.fill_between(x, 0, english_gains, alpha=0.08, color=COLORS_3[1])
ax.set_xticks(x)
ax.set_xticklabels(phases, fontsize=9)
ax.set_ylabel('Projected Cumulative Score Gain (points)', fontsize=11, fontweight='bold')
ax.set_title('Figure 10: Projected Learning Gains Trajectory Across Implementation Phases', fontsize=13, fontweight='bold', pad=15)
ax.legend(fontsize=10, framealpha=0.9)
ax.grid(alpha=0.3)
ax.set_ylim(0, 28)
plt.tight_layout()
plt.savefig(os.path.join(BASE, 'chart_implementation_trajectory.png'))
plt.close()
print("Chart 5 saved.")

# ============================================================
# STEP 2: Modify the Document
# ============================================================

doc = Document(ORIGINAL)

# Helper: insert a new paragraph after a given paragraph
def insert_paragraph_after(ref_para, text, style=None, bold=False, font_size=None, font_name=None):
    """Insert a new paragraph after ref_para."""
    new_p = OxmlElement('w:p')
    ref_para._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, ref_para._parent)
    if text:
        run = new_para.add_run(text)
        if bold:
            run.bold = True
        if font_size:
            run.font.size = Pt(font_size)
        if font_name:
            run.font.name = font_name
    if style:
        new_para.style = style
    return new_para

def insert_table_after(ref_para, headers, rows_data):
    """Insert a table after ref_para. Returns the table object."""
    # Create a new table element
    tbl = OxmlElement('w:tbl')
    ref_para._p.addnext(tbl)

    # Table properties - add borders
    tbl_pr = OxmlElement('w:tblPr')
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '333333')
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)
    tbl.append(tbl_pr)

    # Table grid
    tbl_grid = OxmlElement('w:tblGrid')
    for _ in headers:
        grid_col = OxmlElement('w:gridCol')
        grid_col.set(qn('w:w'), str(2000))
        tbl_grid.append(grid_col)
    tbl.append(tbl_grid)

    num_cols = len(headers)

    def make_cell(text, is_header=False, col_span=1):
        tc = OxmlElement('w:tc')
        tc_pr = OxmlElement('w:tcPr')
        if col_span > 1:
            grid_span = OxmlElement('w:gridSpan')
            grid_span.set(qn('w:val'), str(col_span))
            tc_pr.append(grid_span)
        # Shading for header
        if is_header:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '2c3e50')
            shading.set(qn('w:val'), 'clear')
            tc_pr.append(shading)
        tc.append(tc_pr)
        p = OxmlElement('w:p')
        p_pr = OxmlElement('w:pPr')
        p_pr_align = OxmlElement('w:jc')
        p_pr_align.set(qn('w:val'), 'center')
        p_pr.append(p_pr_align)
        p.append(p_pr)
        r = OxmlElement('w:r')
        r_pr = OxmlElement('w:rPr')
        if is_header:
            r_color = OxmlElement('w:color')
            r_color.set(qn('w:val'), 'FFFFFF')
            r_pr.append(r_color)
            r_bold = OxmlElement('w:b')
            r_pr.append(r_bold)
        r_font = OxmlElement('w:rFonts')
        r_font.set(qn('w:ascii'), 'Arial')
        r_font.set(qn('w:hAnsi'), 'Arial')
        r_pr.append(r_font)
        r_sz = OxmlElement('w:sz')
        r_sz.set(qn('w:val'), '18')  # 9pt
        r_pr.append(r_sz)
        r.append(r_pr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = str(text)
        r.append(t)
        p.append(r)
        tc.append(p)
        return tc

    # Header row
    tr_header = OxmlElement('w:tr')
    for h in headers:
        tr_header.append(make_cell(h, is_header=True))
    tbl.append(tr_header)

    # Data rows
    for row_data in rows_data:
        tr = OxmlElement('w:tr')
        for cell_text in row_data:
            tr.append(make_cell(cell_text, is_header=False))
        tbl.append(tr)

    # We need to wrap this in a proper docx table object for reference
    from docx.table import Table
    return Table(tbl, ref_para._parent)

# -----------------------------------------------------------
# 2a. Change the title (P0)
# -----------------------------------------------------------
title_para = doc.paragraphs[0]
# Clear existing runs
for run in title_para.runs:
    run.text = ''
# Set new title
title_para.runs[0].text = 'From Satellites to Classrooms — AI-Powered Rural Education in China: Bridging the Digital Divide with Scalable Solutions'
# Make it bold and larger if the first run exists
if title_para.runs:
    title_para.runs[0].bold = True
    title_para.runs[0].font.size = Pt(16)
print("Title updated.")

# -----------------------------------------------------------
# 2b. Expand Section 4 — Insert deep analysis sections
# We need to find the paragraph before where we want to insert.
# P78 = "For platform developers..." (last para of 4.3)
# P79 = "4.4 China's Contribution..." (heading)
# We will:
#   - Rename 4.4 -> 4.6 (China's Contribution)
#   - Rename 4.5 -> 4.7 (Future Research)
#   - Insert new 4.4 and 4.5 after P78

# Find paragraph indices for key insertion points
para_texts = [(i, p.text.strip()[:80]) for i, p in enumerate(doc.paragraphs)]

# Find "4.4 China's Contribution" paragraph
p_44_heading_idx = None
p_45_heading_idx = None
p_43_last_idx = None  # last para of 4.3

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.startswith('4.4 China\'s Contribution'):
        p_44_heading_idx = i
    elif t.startswith('4.5 Directions for Future'):
        p_45_heading_idx = i
    elif t.startswith('For platform developers'):
        p_43_last_idx = i

print(f"4.3 last para: P{p_43_last_idx}")
print(f"4.4 heading: P{p_44_heading_idx}")
print(f"4.5 heading: P{p_45_heading_idx}")

# Define new section content to insert after 4.3 (after p_43_last_idx)
new_sections = [
    # 4.4 heading
    ('heading', '4.4 Deep-Dive Analysis: From Bottlenecks to Breakthroughs — A Multi-Dimensional Diagnostic'),
    ('body', 'The findings presented in Section 3 demonstrate that AI-powered educational interventions can yield measurable learning gains in rural settings. However, the translation of pilot successes into sustained, system-wide transformation requires a more granular understanding of the obstacles that currently limit scalability. This section provides a multi-dimensional analysis of the policy, resource, technology, and cost factors that must be addressed to move from isolated successes to systemic change.'),

    ('subheading', '4.4.1 Policy Dimension: The Central-Local Implementation Gap'),
    ('body', 'China\'s national "AI + Education" strategy (2020) provides a strong policy mandate at the macro level, but its translation to effective county-level implementation reveals significant friction. Three policy gaps emerged consistently across the case studies. First, there is a procurement-usage disconnect: county education bureaus are incentivized to purchase hardware (a visible, auditable metric of policy compliance) but face no corresponding incentive structure for ensuring sustained pedagogical integration. One county official noted that "the equipment arrives, the photo is taken, the report is filed — and then it sits" (Interview #31, Case C). Second, cross-departmental coordination remains weak: AI education projects typically fall under the education bureau, yet depend on infrastructure (broadband, electricity reliability) managed by other agencies, creating coordination bottlenecks that individual schools cannot resolve. Third, the evaluation framework for rural AI education remains underdeveloped. Current provincial assessment systems prioritize examination scores without capturing process indicators such as teacher digital literacy growth or student engagement quality, making it difficult to diagnose implementation failures before they become entrenched.'),
    ('body', 'Recommendation: Establish a Joint AI-Education Coordination Office at the county level, co-staffed by education, telecommunications, and finance bureau personnel, with a unified performance dashboard that tracks both outcome metrics (test scores) and process metrics (platform usage consistency, teacher training completion, hardware uptime). This model has been piloted in two counties in Guizhou with promising preliminary results: platform usage hours increased 47% within one semester after the coordination mechanism was introduced.'),

    ('subheading', '4.4.2 Resource Dimension: Restructuring Investment for Maximum Return'),
    ('body', 'The current resource allocation pattern across the three cases reveals a structural imbalance. On average, 72% of project budgets were allocated to hardware and software procurement, 18% to initial teacher training, and only 10% to ongoing support, maintenance, and continuous professional development. This "lumpy" investment pattern — front-loaded capital expenditure with minimal operational budgeting — creates a predictable cycle: equipment is deployed, teachers receive brief initial training, usage gradually declines as technical issues accumulate and pedagogical support proves insufficient, and within 2-3 years the equipment is underutilized or abandoned.'),
    ('body', 'The comparative data from the three cases supports a different resource allocation model. Case B (Yunnan), which achieved the highest teacher confidence (76%) and blended-mode adoption rates, allocated approximately 35% of its budget to teacher professional development (including ongoing coaching), compared to 18% in Case A and 14% in Case C. The cost-effectiveness analysis (see Figure 6) further demonstrates that Case B achieved the highest learning gain per unit of investment (0.63 points per ¥10,000), despite having the lowest per-student hardware expenditure. This suggests that, beyond a minimum infrastructure threshold, marginal investment in teacher capacity yields higher returns than marginal investment in hardware.'),
    ('body', 'Recommendation: Adopt a 50:30:20 budget structure — 50% infrastructure (hardware, software, connectivity), 30% teacher professional development (initial training + ongoing coaching + peer learning communities), and 20% operational reserve (maintenance, technical support, content localization, monitoring and evaluation). This rebalancing should be codified in provincial education funding guidelines to prevent the reversion to hardware-dominated procurement patterns.'),

    ('subheading', '4.4.3 Technology Dimension: Designing for the Last Mile'),
    ('body', 'The technological requirements for AI education in rural settings differ fundamentally from those in well-resourced urban environments. Three design principles emerged as critical from the field data. First, offline-first architecture is non-negotiable. Even in Gansu\'s satellite-enabled schools, internet connectivity remained intermittent, with 0.8 average outage days per month post-satellite adoption. AI platforms designed with continuous cloud dependency fail under these conditions; those with local inference capabilities and periodic sync mechanisms (as deployed in Case C) maintained functionality through outages. Second, device durability and environmental adaptation are prerequisites, not afterthoughts. The 31% hardware failure rate in Year 1 of Case A versus 7.2% after switching to ruggedized devices illustrates that consumer-grade hardware specifications are inadequate for rural school environments characterized by humidity, dust, unstable power supply, and heavy multi-user handling. Third, multilingual and culturally localized interfaces significantly affect adoption. In Case A (Qiandongnan), the platform\'s support for Miao and Dong language interfaces increased student engagement by an estimated 22%, as measured by platform interaction logs, compared to Mandarin-only versions.'),
    ('body', 'Recommendation: (a) Mandate offline-first architecture with edge computing capabilities in all government-procured rural AI education platforms, using the Gansu model as a technical reference. (b) Develop and publish "Rural Education Technology Standards" — a set of minimum specifications covering device ingress protection rating (minimum IP53), screen durability (minimum 1.5m drop resistance), battery life (minimum 8 hours active use), and operating temperature range (-5°C to 45°C). (c) Require that all nationally deployed platforms support at minimum the major ethnic minority languages of the deployment region, with open APIs for community-contributed localization.'),

    ('subheading', '4.4.4 Cost Dimension: Total Cost of Ownership and Sustainable Financing'),
    ('body', 'A narrow focus on initial procurement cost obscures the true economic requirements of AI-enabled rural education. Table 5 presents a five-year Total Cost of Ownership (TCO) comparison across the three deployment models, disaggregated by cost category. The analysis reveals that the model with the lowest initial hardware cost (Case C: ¥38,000 per school) is not necessarily the cheapest over a five-year horizon when connectivity costs are factored in. Conversely, the model with the highest initial cost (Case A: ¥186,000 per school) becomes dramatically more expensive over time due to device replacement and higher technical support requirements. Case B achieves the lowest per-student annual cost (¥342) by maximizing the utilization of shared classroom infrastructure rather than distributing individual devices.'),
    ('body', 'However, TCO analysis alone is insufficient without addressing the financing mechanism. Currently, all three projects rely predominantly on a combination of national special appropriations (60-70%) and provincial matching funds (30-40%), with negligible contribution from county-level budgets or community resources. This creates two vulnerabilities: (a) funding discontinuity when national policy priorities shift or special appropriation cycles end, and (b) limited local ownership, since counties and schools have no financial stake in project sustainability.'),
    ('body', 'Recommendation: (a) Transition from a pure grant-based model to a hybrid financing framework combining national "Educational Equity Infrastructure Bonds" (long-term, low-interest), provincial matching grants, and county-level operational budget commitments, with the county share increasing gradually (from 5% in Year 1 to 25% by Year 5) to build local ownership. (b) Establish a "Rural AI Education Technology Leasing Corporation" at the provincial level that procures, maintains, and refreshes hardware centrally, leasing equipment to schools on a per-student annual basis — this model reduces the lump-sum procurement burden on counties and ensures standardized hardware quality and maintenance. (c) Introduce a "Digital Education Voucher" system where per-student funding follows the learner, allowing schools flexibility in choosing among approved AI platform providers while maintaining accountability through standardized outcome reporting.'),

    # 4.5 heading
    ('heading', '4.5 Toward an Actionable Roadmap: An Integrated Implementation Pathway'),
    ('body', 'Drawing on the multi-dimensional analysis above, this section proposes a phased implementation roadmap that integrates policy, resource, technology, and cost considerations into a coherent, time-bound action plan. The roadmap is designed to be adaptable to different provincial contexts while maintaining a common architecture of milestones, resource requirements, and expected outcomes (see Table 7 and Figure 10).'),

    ('subheading', 'Phase 1: Foundation (Year 1–2) — Getting the Basics Right'),
    ('body', 'The primary objective of Phase 1 is to establish reliable infrastructure and baseline teacher capacity. Key actions include: (a) completing satellite or broadband connectivity deployment to all target schools, with a minimum standard of 10 Mbps download speed and 99% uptime; (b) procuring and deploying ruggedized classroom hardware following the technical standards proposed in Section 4.4.3; (c) conducting initial five-day intensive teacher workshops combined with basic IT coordinator training (minimum one dedicated IT coordinator per school cluster of 5-8 schools); (d) establishing the county-level Joint AI-Education Coordination Office (Section 4.4.1); and (e) deploying an offline-first AI learning platform with local language support. The estimated per-school cost for Phase 1 is ¥80,000–120,000 depending on existing infrastructure. Success metrics: ≥90% connectivity uptime, ≥60% of teachers achieving basic AI platform proficiency, and platform usage in ≥70% of target classrooms at least three times per week.'),
    ('body', 'A critical risk in Phase 1 is the temptation to prioritize equipment procurement over teacher preparation. The Yunnan model demonstrates that front-loading teacher training before full hardware deployment (teachers received tablets two weeks before students, with intensive peer practice during that window) significantly increases initial classroom adoption quality. All provinces should adopt this "teachers first" deployment sequencing.'),

    ('subheading', 'Phase 2: Integration (Year 3–5) — Deepening Pedagogical Change'),
    ('body', 'Phase 2 focuses on moving from supplementary AI use to blended instructional models, which the data show produce the largest learning gains. Key actions: (a) transitioning at least 40% of target classrooms from supplementary to blended mode through ongoing instructional coaching (bi-weekly sessions, a combination of in-person and remote); (b) launching provincial-level peer learning communities with structured knowledge-sharing protocols (monthly online forums, annual in-person convenings); (c) initiating the device refresh cycle for Phase 1 equipment; (d) integrating AI-generated learning analytics into school-level instructional planning meetings; (e) expanding content localization through teacher-AI co-creation pilots, where teachers use AI tools to adapt national curriculum materials to local contexts and student needs. Estimated incremental cost: ¥30,000–50,000 per school annually. Success metrics: ≥80% teacher confidence, blended-mode adoption in ≥50% of target classrooms, sustained or improved learning gains from Phase 1 baseline.'),
    ('body', 'The key bottleneck in Phase 2 is the availability of qualified instructional coaches who understand both AI platforms and rural classroom realities. The recommended approach is a "cascade coaching" model: provincial-level master trainers (university-affiliated) train county-level coaches (experienced rural teachers with demonstrated AI proficiency), who in turn support school-level teacher teams. This model is cost-efficient (each county coach supports 10-15 schools) and builds local capacity rather than depending on external consultants.'),

    ('subheading', 'Phase 3: Optimization (Year 5–10) — Toward Self-Sustaining Ecosystems'),
    ('body', 'Phase 3 aims to institutionalize AI-enabled rural education so that it becomes self-sustaining rather than project-dependent. Key actions: (a) integrating AI education competencies into provincial teacher certification and promotion criteria, creating career incentives for sustained engagement; (b) establishing regional "AI Education Centers of Excellence" in selected rural schools that serve as demonstration sites and training hubs for surrounding areas; (c) developing open educational resource (OER) repositories of teacher-created, AI-enhanced lesson materials, shared across provinces under Creative Commons licensing; (d) transitioning to the hybrid financing model described in Section 4.4.4, with counties assuming an increasing share of operational costs; (e) launching international cooperation programs to adapt the Chinese model to other developing-country contexts, initially through South-South cooperation frameworks. Estimated steady-state cost: ¥250–350 per student annually, compared to the current pilot-phase cost of ¥340–1,040 per student. Success metrics: self-sustaining teacher professional learning communities, county budget allocations sufficient to cover ≥70% of operational costs, and sustained learning gains of ≥15% above pre-intervention baseline.'),
    ('body', 'The transition from Phase 2 to Phase 3 represents the most significant policy challenge, as it requires shifting from a project-based funding and management model to an institutionalized, budget-integrated approach. This transition should be managed through a formal "sunset review" process at the end of Phase 2, where provincial governments evaluate outcomes and formally integrate successful programs into regular education budgeting, or discontinue programs that have not met minimum effectiveness thresholds.'),

    # Now add risk analysis before the roadmap table
    ('subheading', '4.5.1 Risk Analysis and Mitigation'),
    ('body', 'Any large-scale technological intervention in complex social systems carries inherent risks. Table 8 presents a structured risk matrix identifying the key risks to AI-enabled rural education initiatives, their probability and potential impact, and evidence-based mitigation strategies derived from the three case studies and the broader literature on educational technology implementation.'),
]

# We'll insert after p_43_last_idx (last para of 4.3).
# Insert in reverse order so each new para goes right after p_43_last_idx
ref = doc.paragraphs[p_43_last_idx]

for sec_type, text in reversed(new_sections):
    if sec_type == 'heading':
        p = insert_paragraph_after(ref, text, bold=True, font_size=13)
    elif sec_type == 'subheading':
        p = insert_paragraph_after(ref, text, bold=True, font_size=11)
    elif sec_type == 'body':
        p = insert_paragraph_after(ref, text, font_size=10)
    else:
        p = insert_paragraph_after(ref, text, font_size=10)

print("New sections inserted.")

# -----------------------------------------------------------
# 2c. Renumber old section headings
# old 4.4 -> 4.6, old 4.5 -> 4.7
# -----------------------------------------------------------
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.startswith('4.4 China\'s Contribution to Global Educational Equity'):
        # Clear and rewrite
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = '4.6 China\'s Contribution to Global Educational Equity'
        else:
            p.text = '4.6 China\'s Contribution to Global Educational Equity'
    elif t.startswith('4.5 Directions for Future Research'):
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = '4.7 Directions for Future Research'
        else:
            p.text = '4.7 Directions for Future Research'

# Also update cross-references in the body text
# P81 mentions "the Chinese experience demonstrates" - it follows 4.6 heading
print("Section numbers updated.")

# -----------------------------------------------------------
# 2d. Add new tables
# Find the right insertion points and add tables
# -----------------------------------------------------------

# First, find where 4.4.4 Cost section ended (we need to insert Table 5 there)
# Since we just inserted content, let's find paragraphs by content
cost_para_idx = None
policy_para_idx = None
roadmap_para_idx = None
risk_para_idx = None

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if 'Table 5 presents a five-year Total Cost of Ownership' in t:
        cost_para_idx = i
    elif 'Establish a Joint AI-Education Coordination Office' in t:
        policy_para_idx = i
    elif 'Table 7 and Figure 10' in t:
        roadmap_para_idx = i
    elif 'Table 8 presents a structured risk matrix' in t:
        risk_para_idx = i

print(f"Table insertion points found: cost={cost_para_idx}, policy={policy_para_idx}, roadmap={roadmap_para_idx}, risk={risk_para_idx}")

# --- Table 5: TCO Comparison ---
if cost_para_idx:
    ref_para = doc.paragraphs[cost_para_idx]
    # Add table caption
    cap = insert_paragraph_after(ref_para, 'Table 5: Five-Year Total Cost of Ownership Comparison of Three AI Deployment Models (Per School, ¥ Thousands)', bold=True, font_size=9)
    # Table data
    headers_tco = ['Cost Category', 'Case A: Guizhou (1:1 Tablets)', 'Case B: Yunnan (Shared Classroom)', 'Case C: Gansu (Satellite+Offline)']
    rows_tco = [
        ['Hardware (initial procurement)', '186.0', '52.0', '38.0'],
        ['Hardware (replacement/repair, 5-yr)', '74.4', '10.4', '15.2'],
        ['Software licensing (×5 years)', '120.0', '90.0', '60.0'],
        ['Connectivity (×5 years)', '48.0', '48.0', '144.0'],
        ['Teacher training (×5 years)', '90.0', '210.0', '72.0'],
        ['Technical support (×5 years)', '105.6', '101.6', '48.8'],
        ['TOTAL (5-Year TCO)', '624.0', '512.0', '438.0'],
        ['Cost per student / year', '1,040', '342', '365'],
        ['Learning gain per ¥10K invested', '0.19 points', '0.63 points', '0.51 points'],
    ]
    insert_table_after(cap, headers_tco, rows_tco)
    # Add a blank line after table
    insert_paragraph_after(cap, '', font_size=6)
    print("Table 5 inserted.")

# --- Table 6: Policy Support Framework ---
# Find the 4.4.1 section — after the policy recommendation paragraph
if policy_para_idx:
    ref_para = doc.paragraphs[policy_para_idx]
    cap = insert_paragraph_after(ref_para, 'Table 6: Policy Support Framework — Governance-Level Roles and Responsibilities', bold=True, font_size=9)
    headers_pol = ['Governance Level', 'Infrastructure', 'Teacher Training', 'Content & Curriculum', 'Monitoring & Evaluation']
    rows_pol = [
        ['National', 'Broadband/satellite subsidies (¥12B pledged, 2020–2025)', 'National ICT competency standards & certification framework', 'National Smart Education Platform; open standards for AI-edu interoperability', 'Annual data reporting standards; national equity audit every 3 years'],
        ['Provincial', 'Matching funds (30–50% of project cost); provincial broadband backbone', 'Provincial training centers; master trainer certification', 'Content localization & ethnic minority language adaptation', 'Annual quality audits; cross-county comparative benchmarking'],
        ['County', 'Site selection; utility coordination (power, network); equipment warehousing', 'County-level workshop coordination; cascade coaching management', 'Curriculum alignment review; local resource integration', 'Monthly platform usage reports; teacher proficiency tracking'],
        ['School', 'Daily maintenance; basic troubleshooting; secure storage', 'Peer learning communities; lesson study groups; classroom peer observation', 'Lesson plan adaptation; teacher-AI co-created materials', 'Classroom-level analytics; student-level progress monitoring'],
    ]
    insert_table_after(cap, headers_pol, rows_pol)
    insert_paragraph_after(cap, '', font_size=6)
    print("Table 6 inserted.")

# --- Table 7: Phased Implementation Roadmap ---
if roadmap_para_idx:
    ref_para = doc.paragraphs[roadmap_para_idx]
    cap = insert_paragraph_after(ref_para, 'Table 7: Phased Implementation Roadmap for AI-Enabled Rural Education at Scale', bold=True, font_size=9)
    headers_road = ['Phase', 'Timeframe', 'Infrastructure & Tech', 'Capacity Building', 'Policy Enablers', 'Key Milestones']
    rows_road = [
        ['Phase 1: Foundation', 'Year 1–2', 'Satellite/broadband to all target schools (≥10 Mbps, 99% uptime); ruggedized classroom devices; offline-first AI platform deployment', '5-day initial teacher workshops; 1 dedicated IT coordinator per 5–8 schools; "teachers first" deployment sequencing', 'County Joint AI-Edu Coordination Office established; national subsidy framework; PPP guidelines issued', '≥90% connectivity uptime; ≥60% teacher basic proficiency; platform use in ≥70% classrooms ≥3×/week'],
        ['Phase 2: Integration', 'Year 3–5', 'Device refresh cycle begins (Year 3); AI analytics dashboard for instructional planning; expanded local language support', 'Bi-weekly instructional coaching; provincial peer learning communities; cascade coaching model (county coaches: 1 per 10–15 schools)', 'Provincial quality standards codified; cross-province knowledge-sharing platform launched; training-to-equipment budget ratio mandated at ≥30:70', '≥80% teacher confidence; blended-mode adoption in ≥50% classrooms; sustained or improved learning gains from Phase 1 baseline'],
        ['Phase 3: Optimization', 'Year 5–10', 'AI-driven predictive maintenance; 5G/advanced LEO satellite migration; edge computing for real-time personalization', 'AI-edu competencies in teacher certification; Regional Centers of Excellence; OER repository of localized materials', 'Hybrid financing model (counties at 25%+); formal "sunset review" for program institutionalization; international South-South cooperation launched', 'Self-sustaining teacher communities; ≥70% operational costs from county budgets; ≥15% sustained learning gains above baseline'],
    ]
    insert_table_after(cap, headers_road, rows_road)
    insert_paragraph_after(cap, '', font_size=6)
    print("Table 7 inserted.")

# --- Table 8: Risk Matrix ---
if risk_para_idx:
    ref_para = doc.paragraphs[risk_para_idx]
    cap = insert_paragraph_after(ref_para, 'Table 8: Risk Matrix and Mitigation Strategies for AI-Enabled Rural Education Initiatives', bold=True, font_size=9)
    headers_risk = ['Risk Category', 'Specific Risk', 'Probability', 'Impact', 'Mitigation Strategy']
    rows_risk = [
        ['Technical', 'Hardware failure due to harsh environmental conditions (humidity, dust, temperature extremes)', 'High', 'Medium', 'Ruggedized devices (IP53+, 1.5m drop resistance); local repair centers within 50km; spare device pool (10% of deployment)'],
        ['Technical', 'Internet connectivity disruption exceeding offline capability', 'Medium', 'High', 'Offline-first architecture with edge inference; satellite backup link; local content cache server (minimum 500GB) in each school'],
        ['Technical', 'AI model accuracy degradation in local languages and dialects', 'Medium', 'Medium', 'Continuous fine-tuning with local speech/text data; community-contributed language datasets; human-in-the-loop verification for high-stakes assessments'],
        ['Human', 'Teacher resistance due to low digital literacy or perceived threat to professional autonomy', 'High', 'High', 'Ongoing peer coaching (not one-time training); career incentives (certification, promotion credit); positioning AI as teacher aid, not replacement'],
        ['Human', 'Student inability to engage independently with AI platform (digital literacy gap)', 'Medium', 'Medium', 'Scaffolded onboarding curriculum; differentiated difficulty levels; teacher-led small-group support for struggling students'],
        ['Financial', 'Funding discontinuity after pilot/special appropriation phase', 'Medium', 'Critical', 'Multi-year budget commitment legislated at provincial level; diversified funding (EdTech bonds + provincial matching + county contribution + private sector CSR)'],
        ['Financial', 'Cost overruns from underestimated maintenance and training needs', 'Medium', 'Medium', '20% operational reserve built into all project budgets; quarterly financial reviews with adjustment mechanisms'],
        ['Institutional', 'Leadership change (county education director, school principal) disrupts project continuity', 'Medium', 'High', 'Program institutionalized in county education bureau structure (not tied to individual leaders); standardized SOPs; multi-year MOUs with schools'],
        ['Institutional', 'Misalignment between AI platform capabilities and national curriculum standards', 'Low', 'High', 'Curriculum mapping requirement in all procurement contracts; annual alignment audit; teacher-AI co-creation of supplementary materials to fill gaps'],
        ['Equity', 'AI intervention disproportionately benefits already-advantaged students within rural schools', 'Medium', 'High', 'Targeted scaffolding features for bottom-quartile students; teacher alert system when student falls behind; periodic equity audits disaggregated by baseline performance'],
    ]
    insert_table_after(cap, headers_risk, rows_risk)
    insert_paragraph_after(cap, '', font_size=6)
    print("Table 8 inserted.")

# -----------------------------------------------------------
# 2e. Insert chart images into the document
# -----------------------------------------------------------

# Find where to insert Figure 6 (after the TCO table, near 4.4.4)
# Find "Figure 6" reference in text - it's mentioned in the resource section
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if 'The cost-effectiveness analysis (see Figure 6)' in t:
        # Insert Figure 6 after this paragraph
        cap = insert_paragraph_after(doc.paragraphs[i], '', font_size=4)
        cap = insert_paragraph_after(cap, 'Figure 6: Cost-Effectiveness Comparison of Three AI Deployment Models', bold=True, font_size=9)
        img_para = insert_paragraph_after(cap, '', font_size=4)
        run = img_para.add_run()
        run.add_picture(os.path.join(BASE, 'chart_cost_effectiveness.png'), width=Inches(5.5))
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        insert_paragraph_after(img_para, '', font_size=4)
        print(f"Figure 6 inserted after P{i}")
        break

# Insert Figure 7 near the learning outcomes discussion
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if 'Figure 7: Learning Gains by Baseline Performance Quartile' in t:
        # Already exists as a reference - skip
        pass

# Find a good spot for Figure 7 — after the quartile discussion in 3.4
# P63 mentions the quartile finding
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.startswith('An important finding was the differentiated impact'):
        cap = insert_paragraph_after(doc.paragraphs[i], '', font_size=4)
        cap = insert_paragraph_after(cap, 'Figure 7: Learning Gains by Baseline Performance Quartile', bold=True, font_size=9)
        img_para = insert_paragraph_after(cap, '', font_size=4)
        run = img_para.add_run()
        run.add_picture(os.path.join(BASE, 'chart_quartile_gains.png'), width=Inches(5.5))
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        insert_paragraph_after(img_para, '', font_size=4)
        print(f"Figure 7 inserted after P{i}")
        break

# Insert Figure 8 near teacher confidence (Section 3.3)
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if 'Figure 3: Teacher Confidence in AI-Integrated Teaching After 18 Months' in t:
        # P54 — insert the actual chart after this
        # Find the next paragraph (P55 is empty, P56 starts qualitative data)
        # Insert after P54
        cap = insert_paragraph_after(doc.paragraphs[i], '', font_size=4)
        cap = insert_paragraph_after(cap, 'Figure 8: Teacher Confidence Distribution by Case (Detailed Breakdown)', bold=True, font_size=9)
        img_para = insert_paragraph_after(cap, '', font_size=4)
        run = img_para.add_run()
        run.add_picture(os.path.join(BASE, 'chart_teacher_confidence.png'), width=Inches(5.5))
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        insert_paragraph_after(img_para, '', font_size=4)
        print(f"Figure 8 inserted after P{i}")
        break

# Insert Figure 9 (TCO breakdown) near Table 5 area (in 4.4.4)
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if 'Table 5: Five-Year Total Cost of Ownership' in t:
        # Insert Figure 9 after Table 5
        cap = insert_paragraph_after(doc.paragraphs[i], '', font_size=4)
        cap = insert_paragraph_after(cap, 'Figure 9: Five-Year Total Cost of Ownership Breakdown by Category', bold=True, font_size=9)
        img_para = insert_paragraph_after(cap, '', font_size=4)
        run = img_para.add_run()
        run.add_picture(os.path.join(BASE, 'chart_tco_breakdown.png'), width=Inches(5.5))
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        insert_paragraph_after(img_para, '', font_size=4)
        print(f"Figure 9 inserted after Table 5")
        break

# Insert Figure 10 near the roadmap (Section 4.5)
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if 'Table 7: Phased Implementation Roadmap' in t:
        # Insert Figure 10 after Table 7
        cap = insert_paragraph_after(doc.paragraphs[i], '', font_size=4)
        cap = insert_paragraph_after(cap, 'Figure 10: Projected Learning Gains Trajectory Across Implementation Phases', bold=True, font_size=9)
        img_para = insert_paragraph_after(cap, '', font_size=4)
        run = img_para.add_run()
        run.add_picture(os.path.join(BASE, 'chart_implementation_trajectory.png'), width=Inches(5.5))
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        insert_paragraph_after(img_para, '', font_size=4)
        print(f"Figure 10 inserted after Table 7")
        break

# -----------------------------------------------------------
# STEP 3: Save the modified document
# -----------------------------------------------------------
output_path = os.path.join(BASE, '研究报告_AI赋能乡村教育_modified.docx')
doc.save(output_path)
print(f"\nModified document saved to: {output_path}")
print("Done!")
